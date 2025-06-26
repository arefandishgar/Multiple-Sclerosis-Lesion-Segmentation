import os
import sys
import torch
import torch.nn as nn
import torch.nn.functional as F
import nibabel as nib
import numpy as np
from tqdm import tqdm
import matplotlib.pyplot as plt

class Logger:
    def __init__(self, filename):
        self.terminal = sys.stdout
        self.log = open(filename, "w")

    def write(self, message):
        self.terminal.write(message)
        self.log.write(message)
        self.flush()

    def flush(self):
        self.terminal.flush()
        self.log.flush()

    def close(self):
        self.log.close()

class R2CL(nn.Module):
    def __init__(self, in_channels, out_channels):
        super(R2CL, self).__init__()
        self.conv1 = nn.Conv3d(in_channels, out_channels, kernel_size=3, padding=1)
        self.bn1 = nn.BatchNorm3d(out_channels)
        self.conv2 = nn.Conv3d(out_channels, out_channels, kernel_size=3, padding=1)
        self.bn2 = nn.BatchNorm3d(out_channels)
        self.relu = nn.ReLU(inplace=True)
        self.residual_conv = nn.Conv3d(in_channels, out_channels, kernel_size=1)

    def forward(self, x):
        residual = self.residual_conv(x)
        out = self.relu(self.bn1(self.conv1(x)))
        out = self.bn2(self.conv2(out))
        out += residual
        out = self.relu(out)
        return out

class AttentionGate(nn.Module):
    def __init__(self, F_g, F_l, F_int):
        super(AttentionGate, self).__init__()
        self.W_g = nn.Sequential(
            nn.Conv3d(F_g, F_int, kernel_size=1, stride=1, padding=0, bias=True),
            nn.BatchNorm3d(F_int)
        )
        self.W_x = nn.Sequential(
            nn.Conv3d(F_l, F_int, kernel_size=1, stride=1, padding=0, bias=True),
            nn.BatchNorm3d(F_int)
        )
        self.psi = nn.Sequential(
            nn.Conv3d(F_int, 1, kernel_size=1, stride=1, padding=0, bias=True),
            nn.BatchNorm3d(1),
            nn.Sigmoid()
        )
        self.relu = nn.ReLU(inplace=True)

    def forward(self, g, x):
        g1 = self.W_g(g)
        x1 = self.W_x(x)
        psi = self.relu(g1 + x1)
        psi = self.psi(psi)
        return x * psi

class R2AUNet(nn.Module):
    def __init__(self, in_channels, out_channels):
        super(R2AUNet, self).__init__()
        features = [32, 64, 128, 256]
        
        self.encoder = nn.ModuleList()
        self.pool = nn.MaxPool3d(kernel_size=2, stride=2)
        for i in range(4):
            if i == 0:
                self.encoder.append(R2CL(in_channels, features[i]))
            else:
                self.encoder.append(R2CL(features[i-1], features[i]))
        
        self.bridge = R2CL(features[-1], features[-1]*2)
        
        self.decoder = nn.ModuleList()
        self.attention_gates = nn.ModuleList()
        self.up_conv = nn.ModuleList()
        for i in range(4):
            self.up_conv.append(nn.ConvTranspose3d(features[-i-1]*2, features[-i-1], kernel_size=2, stride=2))
            self.attention_gates.append(AttentionGate(F_g=features[-i-1], F_l=features[-i-1], F_int=features[-i-1]//2))
            self.decoder.append(R2CL(features[-i-1]*2, features[-i-1]))
        
        self.final_conv = nn.Conv3d(features[0], out_channels, kernel_size=1)

    def forward(self, x):
        encoder_features = []
        for enc in self.encoder:
            x = enc(x)
            encoder_features.append(x)
            x = self.pool(x)
        
        x = self.bridge(x)
        
        for i in range(4):
            x = self.up_conv[i](x)
            enc_feature = encoder_features[-i-1]
            
            if x.size()[2:] != enc_feature.size()[2:]:
                x = F.interpolate(x, size=enc_feature.size()[2:], mode='trilinear', align_corners=False)
            
            attention = self.attention_gates[i](g=x, x=enc_feature)
            x = torch.cat([x, attention], dim=1)
            x = self.decoder[i](x)
        
        x = self.final_conv(x)
        return x

def predict_full_volume(model, flair_volume, device, patch_size=(64, 64, 64), stride=32):
    model.eval()
    with torch.no_grad():
        pred_volume = np.zeros_like(flair_volume)
        count_volume = np.zeros_like(flair_volume)
        
        for x in tqdm(range(0, flair_volume.shape[0] - patch_size[0] + 1, stride), desc="Predicting"):
            for y in range(0, flair_volume.shape[1] - patch_size[1] + 1, stride):
                for z in range(0, flair_volume.shape[2] - patch_size[2] + 1, stride):
                    patch = flair_volume[x:x+patch_size[0], y:y+patch_size[1], z:z+patch_size[2]]
                    patch_tensor = torch.from_numpy(patch).float().unsqueeze(0).unsqueeze(0).to(device)
                    
                    output = model(patch_tensor)
                    pred_patch = torch.sigmoid(output).squeeze().cpu().numpy()
                    
                    pred_volume[x:x+patch_size[0], y:y+patch_size[1], z:z+patch_size[2]] += pred_patch
                    count_volume[x:x+patch_size[0], y:y+patch_size[1], z:z+patch_size[2]] += 1
        
        # Average overlapping predictions
        pred_volume = np.divide(pred_volume, count_volume, where=count_volume!=0)
        
    return pred_volume

def threshold_predictions(pred_volume, threshold=0.5):
    """Convert continuous predictions to binary mask."""
    return (pred_volume > threshold).astype(np.float32)

def plot_full_volume(flair, pred_volume, save_path):
    fig, axes = plt.subplots(3, 2, figsize=(15, 20))
    views = ['Axial', 'Coronal', 'Sagittal']
    mid_slices = [flair.shape[i] // 2 for i in range(3)]

    for i, view in enumerate(views):
        if view == 'Axial':
            sl = lambda x: x[mid_slices[i], :, :]
        elif view == 'Coronal':
            sl = lambda x: x[:, mid_slices[i], :]
        else:  # Sagittal
            sl = lambda x: x[:, :, mid_slices[i]]

        axes[i, 0].imshow(sl(flair), cmap='gray')
        axes[i, 0].set_title(f'{view} - Original MRI')
        
        axes[i, 1].imshow(sl(flair), cmap='gray')
        axes[i, 1].imshow(sl(pred_volume), cmap='hot', alpha=0.5)
        axes[i, 1].set_title(f'{view} - Binary Prediction Overlay')

    for ax in axes.ravel():
        ax.axis('off')

    plt.tight_layout()
    plt.savefig(save_path, dpi=300, bbox_inches='tight')
    plt.close()
    print(f"Full volume prediction saved to {save_path}")

def load_and_preprocess_mri(file_path):
    """Load and preprocess MRI data."""
    mri = nib.load(file_path)
    data = mri.get_fdata()
     
    # Normalize the data
    #data = (data - data.min()) / (data.max() - data.min())
    
    return data, mri.affine

def save_prediction_mask(prediction, affine, output_path):
    """Save the prediction mask as a NIfTI file."""
    nib.save(nib.Nifti1Image(prediction.astype(np.float32), affine), output_path)

from docx import Document  # Import for Word document creation

def calculate_metrics(y_true, y_pred):
    y_true = y_true.flatten()
    y_pred = y_pred.flatten()
    
    tp = np.sum((y_true == 1) & (y_pred == 1))
    tn = np.sum((y_true == 0) & (y_pred == 0))
    fp = np.sum((y_true == 0) & (y_pred == 1))
    fn = np.sum((y_true == 1) & (y_pred == 0))
    
    epsilon = 1e-7
    accuracy = (tp + tn) / (tp + tn + fp + fn + epsilon)
    sensitivity = tp / (tp + fn + epsilon)
    specificity = tn / (tn + fp + epsilon)
    precision = tp / (tp + fp + epsilon)
    f1_score = 2 * (precision * sensitivity) / (precision + sensitivity + epsilon)
    dice_score = 2 * tp / (2 * tp + fp + fn + epsilon)
    
    return accuracy, sensitivity, specificity, precision, f1_score, dice_score


def process_main_folder_with_general_metrics(main_folder, output_folder, model, device, patch_size=(64, 64, 64), stride=32, metrics_file="test_metrics.docx", general_metrics_file="general_metrics.docx"):
    """Process a main folder with subfolders containing FLAIR and mask files and compute general metrics."""
    os.makedirs(output_folder, exist_ok=True)

    # Create a Word document to store metrics for each file
    document = Document()
    document.add_heading("Test Data Metrics", level=1)
    table = document.add_table(rows=1, cols=8)
    header_cells = table.rows[0].cells
    header_cells[0].text = "Subfolder"
    header_cells[1].text = "Filename"
    header_cells[2].text = "Accuracy"
    header_cells[3].text = "Sensitivity"
    header_cells[4].text = "Specificity"
    header_cells[5].text = "Precision"
    header_cells[6].text = "F1-Score"
    header_cells[7].text = "Dice Score"

    # Initialize lists to collect metrics for general evaluation
    all_metrics = []

    for subfolder in os.listdir(main_folder):
        subfolder_path = os.path.join(main_folder, subfolder)
        if not os.path.isdir(subfolder_path):
            continue

        # Locate FLAIR and mask files in the subfolder
        flair_file = os.path.join(subfolder_path, "FLAIR.nii.gz")
        mask_file = os.path.join(subfolder_path, "FLAIR_Segmentation.nii.gz")

        if not os.path.exists(flair_file) or not os.path.exists(mask_file):
            print(f"Skipping {subfolder}: Missing FLAIR or mask file.")
            continue

        print(f"Processing subfolder: {subfolder}")
        
        # Load and preprocess FLAIR and mask
        mri_data, affine = load_and_preprocess_mri(flair_file)
        true_mask = nib.load(mask_file).get_fdata()

        # Predict on the full volume
        pred_volume = predict_full_volume(model, mri_data, device, patch_size, stride)
        binary_mask = threshold_predictions(pred_volume)

        # Calculate metrics
        try:
            metrics = calculate_metrics(true_mask, binary_mask)
            all_metrics.append(metrics)  # Collect metrics for general evaluation
            accuracy, sensitivity, specificity, precision, f1_score, dice_score = metrics
        except Exception as e:
            print(f"Error calculating metrics for {flair_file}: {e}")
            accuracy = sensitivity = specificity = precision = f1_score = dice_score = None

        # Save metrics in the file-specific table
        row_cells = table.add_row().cells
        row_cells[0].text = subfolder
        row_cells[1].text = "FLAIR.nii.gz"
        row_cells[2].text = f"{accuracy:.4f}" if accuracy is not None else "N/A"
        row_cells[3].text = f"{sensitivity:.4f}" if sensitivity is not None else "N/A"
        row_cells[4].text = f"{specificity:.4f}" if specificity is not None else "N/A"
        row_cells[5].text = f"{precision:.4f}" if precision is not None else "N/A"
        row_cells[6].text = f"{f1_score:.4f}" if f1_score is not None else "N/A"
        row_cells[7].text = f"{dice_score:.4f}" if dice_score is not None else "N/A"

        # Save the binary prediction mask and plot
        output_subfolder = os.path.join(output_folder, subfolder)
        os.makedirs(output_subfolder, exist_ok=True)
        pred_mask_path = os.path.join(output_subfolder, "pred_mask.nii.gz")
        plot_path = os.path.join(output_subfolder, "plot.png")

        save_prediction_mask(binary_mask, affine, pred_mask_path)
        plot_full_volume(mri_data, binary_mask, plot_path)

        print(f"Saved prediction mask to: {pred_mask_path}")
        print(f"Saved overlay plot to: {plot_path}")

    # Save the Word document for per-file metrics
    document.save(metrics_file)
    print(f"Per-file metrics saved to {metrics_file}")

    # Compute general metrics for the whole dataset
    if all_metrics:
        # Convert list of metrics to a NumPy array for easier mean calculation
        all_metrics = np.array(all_metrics)
        general_metrics = np.mean(all_metrics, axis=0)  # Calculate mean of all metrics
        general_accuracy, general_sensitivity, general_specificity, general_precision, general_f1_score, general_dice_score = general_metrics

        # Save general metrics in a separate Word document
        general_doc = Document()
        general_doc.add_heading("General Metrics for Whole Dataset", level=1)
        general_table = general_doc.add_table(rows=1, cols=7)
        general_header = general_table.rows[0].cells
        general_header[0].text = "Accuracy"
        general_header[1].text = "Sensitivity"
        general_header[2].text = "Specificity"
        general_header[3].text = "Precision"
        general_header[4].text = "F1-Score"
        general_header[5].text = "Dice Score"

        # Add a single row for general metrics
        row_cells = general_table.add_row().cells
        row_cells[0].text = f"{general_accuracy:.4f}"
        row_cells[1].text = f"{general_sensitivity:.4f}"
        row_cells[2].text = f"{general_specificity:.4f}"
        row_cells[3].text = f"{general_precision:.4f}"
        row_cells[4].text = f"{general_f1_score:.4f}"
        row_cells[5].text = f"{general_dice_score:.4f}"

        general_doc.save(general_metrics_file)
        print(f"General metrics saved to {general_metrics_file}")
    else:
        print("No metrics were calculated. Check your data or prediction process.")

def main():
    # Set up logging
    log_dir = "logs"
    os.makedirs(log_dir, exist_ok=True)
    sys.stdout = Logger(os.path.join(log_dir, 'stdout.log'))
    sys.stderr = Logger(os.path.join(log_dir, 'stderr.log'))

    # Set up device
    device = torch.device("cuda" if torch.cuda.is_available() else "cpu")
    print(f"Using device: {device}")
    
    # Load the trained model
    
    #model.load_state_dict(torch.load('r2au_net_model.pth', map_location=device))
    #model.load_state_dict(torch.load('r2au_net_model.pth', map_location=device))
    #model.eval()
    state_dict = torch.load('Second results from the training on vast.ai/best_dice_score_model.pth', map_location='cuda:0')
    new_state_dict = {}
    for key, value in state_dict.items():
        new_key = key.replace("module.", "")  # Remove 'module.' prefix
        new_state_dict[new_key] = value
        
    model = R2AUNet(in_channels=1, out_channels=1)
    model.load_state_dict(new_state_dict)
    device = torch.device('cuda:0')
    model = model.to(device)

    # Set up input and output folders
    main_folder = r"/home/aref/Desktop/Brain MRI - MS Project/Codes for patchy model/Entire Data/Test"
    output_folder = r"/home/aref/Desktop/Brain MRI - MS Project/Codes for patchy model/New folder 2"  

    # Process the main folder and calculate metrics
    process_main_folder_with_general_metrics(
        main_folder, output_folder, model, device,
        metrics_file="metrics_tests.docx",
        general_metrics_file="metrics_general.docx"
    )

if __name__ == "__main__":
    try:
        main()
    except Exception as e:
        print(f"An error occurred: {str(e)}")
    finally:
        # Close the log files
        if isinstance(sys.stdout, Logger):
            sys.stdout.close()
        if isinstance(sys.stderr, Logger):
            sys.stderr.close()
        # Restore original stdout and stderr
        sys.stdout = sys.__stdout__
        sys.stderr = sys.__stderr__



